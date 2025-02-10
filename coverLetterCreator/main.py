print('\n\n')
import openai
import os
api_key=os.getenv("OPEN_API_KEY")
openai.api_key = api_key
from docx import Document
from docx.shared import Pt


class gpt:

    def __init__(self, resume, jobDescription, additionalNotes):
        self.resume = resume
        self.jobDescription = jobDescription
        self.additionalNotes = additionalNotes        
        
    def testingForFlask(self):
        print("testingForFlask method called!")
        print(f"length of field1: {len(self.resume)}\nlen field2: {len(self.jobDescription)}\nlen field3: {len(self.additionalNotes)}")


    def talkToGPT(self):
        if len(self.additionalNotes) > 3:
            prompt = f"This is my resume: \"{self.resume}\". This is the job description: \"{self.jobDescription}\". {self.additionalNotes}. Write a cover letter for me."
        else:
            prompt = f"This is my resume: \"{self.resume}\". This is the job description: <\"{self.jobDescription}\">. Write a cover letter for me."
        print(prompt)
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            #model="4o",
            messages=[{"role":"user","content": prompt}]
        )
        response = response['choices'][0]['message']['content']
        #print(f"GPT SAID: {response}")
        return response
    
    def createDocxFile(self, paragraphs: list):
        
        # Create a new document
        doc = Document()
        # Add Name (Big and Bold)
        name = doc.add_paragraph()
        name_run = name.add_run("Your Name")
        name_run.bold = True
        name_run.font.size = Pt(20)

        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

        doc.save("Cover_Letter_By_AI.docx")
        print("Cover letter created successfully!")


if __name__== "__main__":

    instance1 = gpt("John Doe", "turquoise", "")
    instance1.createDocxFile()